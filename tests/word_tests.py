#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @Time    : 2020/10/14 9:33
# @Author  : 李加林
# @FileName: word_tests.py
# @Software: PyCharm
# @Blog    ：https://blog.csdn.net/weixin_43972976
# @github : https://github.com/karinlee1988/
# @gitee : https://gitee.com/karinlee/

import docx
from docx import Document  # 导入库
# from workoverflow.fileio import record_csv
# 读取word中的表格并生成csv
# path = 'source.docx'  # 文件路径
# document = Document(path)  # 读入文件
# tables = document.tables  # 获取文件中的表格集
# print(tables)
# for table in tables:
#     for i, row in enumerate(table.rows):  # 读每行
#         row_content = []
#         for cell in row.cells:  # 读一行中的所有单元格
#             c = cell.text
#             row_content.append(c)
#         record_csv(row_content,"result1.csv")

        # print (row_content) #以列表形式导出每一行数据

def replace_text2(doc,old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    text = i.text.replace(old_text, new_text)
                    i.text = text
def replace_text(doc:[Document],old_text:str, new_text:str) -> None:
    """
    将某个文档的某个词进行替换

    :param doc:   要处理的文件
    :type doc: Document

    :param old_text:  旧的词
    :type old_text: str

    :param new_text: 新的词
    :type new_text: str

    :return: None
    """
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    text = i.text.replace(old_text, new_text)
                    i.text = text


docu = Document('追回决定书模板.docx')
replace_text(docu,"{姓名}","周有德")
replace_text(docu,"{身份证号码}","441881100000000000")
docu.save("result.docx")